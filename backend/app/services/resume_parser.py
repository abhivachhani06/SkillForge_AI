import fitz  # PyMuPDF
import docx
import io
import logging

logger = logging.getLogger("resume_parser")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise ValueError(f"Failed to parse PDF file: {str(e)}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    text = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
        # Also parse tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise ValueError(f"Failed to parse Word document: {str(e)}")
    return "\n".join(text)

def extract_text(file_bytes: bytes, file_name: str) -> str:
    """Identifies the file type and extracts the raw text content."""
    lower_name = file_name.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are allowed.")
